import fitz  # PyMuPDF
from docx import Document


def extract_text(file_path: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT files."""
    if file_path.lower().endswith(".pdf"):
        return _extract_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return _extract_docx(file_path)
    else:
        return _extract_txt(file_path)


def _extract_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages = []
    for page in doc:
        text = page.get_text("text")
        pages.append(text)
    doc.close()
    return "\n".join(pages)


def _extract_docx(path: str) -> str:
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()